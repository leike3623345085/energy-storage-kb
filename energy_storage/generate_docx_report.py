#!/usr/bin/env python3
"""
生成述责述廉报告 Word 文档
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
from pathlib import Path

def set_chinese_font(run, font_name='宋体', size=12, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(size)
    font.bold = bold

def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    para = doc.add_heading(level=level)
    run = para.add_run(text)
    if level == 1:
        set_chinese_font(run, '黑体', 16, True)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        set_chinese_font(run, '黑体', 14, True)
    else:
        set_chinese_font(run, '黑体', 12, True)
    return para

def add_paragraph_zh(doc, text, bold=False, indent=True):
    """添加中文段落"""
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.first_line_indent = Inches(0.5)
    run = para.add_run(text)
    set_chinese_font(run, '仿宋_GB2312' if not bold else '黑体', 12, bold)
    para.paragraph_format.line_spacing = 1.5
    return para

def generate_report(output_path=None):
    """生成述责述廉报告"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '仿宋_GB2312'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    style.font.size = Pt(12)
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('述责述廉报告')
    set_chinese_font(run, '方正小标宋简体', 22, True)
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('党务工作部信息技术工程师')
    set_chinese_font(run, '楷体_GB2312', 14)
    
    doc.add_paragraph()  # 空行
    
    # 一、履行党风廉政建设"一岗双责"情况
    add_heading_zh(doc, '一、履行党风廉政建设"一岗双责"情况', level=2)
    
    add_heading_zh(doc, '（一）强化政治理论学习', level=3)
    add_paragraph_zh(doc, '全年参加支部集中学习12次，自学《中国共产党纪律处分条例》《中国共产党廉洁自律准则》等党纪法规，撰写心得体会4篇。通过学习，进一步筑牢拒腐防变的思想防线，增强"四个意识"、坚定"四个自信"、做到"两个维护"。')
    
    add_heading_zh(doc, '（二）落实岗位廉政责任', level=3)
    add_paragraph_zh(doc, '在信息技术日常工作中，严格执行以下廉政要求：')
    
    add_paragraph_zh(doc, '1. 采购环节：涉及软硬件采购时，严格执行集中采购制度，主动回避与供应商私下接触，所有技术参数论证均邀请纪检人员列席，确保公开透明。')
    
    add_paragraph_zh(doc, '2. 数据管理：严格遵守客户信息保密规定，未发生违规查询、下载、传输敏感数据行为。全年配合完成数据安全自查3次，未发现违规问题。')
    
    add_paragraph_zh(doc, '3. 项目实施：参与系统建设过程中，严格执行验收标准，不收受开发商礼金礼品，不参加可能影响公正执行公务的宴请活动。')
    
    add_heading_zh(doc, '（三）配合监督执纪工作', level=3)
    add_paragraph_zh(doc, '全年配合部门完成廉政风险排查2次，针对信息技术岗位梳理出"系统权限分配""外包人员管理""设备报废处置"等3项廉政风险点，并制定相应防控措施。')
    
    # 二、遵守党章党规党纪情况
    add_heading_zh(doc, '二、遵守党章党规党纪、廉洁从业及执行中央八项规定精神情况', level=2)
    
    add_heading_zh(doc, '（一）个人廉洁自律情况', level=3)
    add_paragraph_zh(doc, '1. 无违纪违规行为：本年度未发生违反党章党规党纪、廉洁从业要求的行为，未受到党纪政纪处分或组织处理。')
    
    add_paragraph_zh(doc, '2. 严格执行八项规定：未接受管理服务对象安排的宴请、旅游、健身等活动；未违规收送礼品礼金、消费卡券；未违规使用公务用车、办公用房；未违规操办婚丧喜庆事宜。')
    
    add_paragraph_zh(doc, '3. 如实报告个人事项：按时如实填报个人有关事项报告表，无漏报瞒报。')
    
    add_heading_zh(doc, '（二）家风家教情况', level=3)
    add_paragraph_zh(doc, '注重家庭家教家风建设，配偶及子女无经商办企业行为，无利用本人职权或职务影响谋取私利情况。')
    
    # 三、下一步努力方向
    add_heading_zh(doc, '三、下一步努力和改进方向', level=2)
    
    add_heading_zh(doc, '（一）存在问题', level=3)
    add_paragraph_zh(doc, '1. 理论学习深度不够：对党纪法规的学习还停留在通读层面，结合信息技术岗位实际的思考不够深入。')
    add_paragraph_zh(doc, '2. 风险防控意识需加强：对外包人员日常管理的监督还不够细致，存在"老好人"思想。')
    add_paragraph_zh(doc, '3. 廉政教育形式单一：在科室内部廉政提醒主要靠会议传达，创新手段不多。')
    
    add_heading_zh(doc, '（二）改进措施', level=3)
    add_paragraph_zh(doc, '1. 深化理论武装：制定个人学习计划，每月至少研读1篇党风廉政建设相关文件或案例，结合岗位实际撰写剖析材料。')
    add_paragraph_zh(doc, '2. 织密监督网络：建立外包人员廉政档案，实行季度谈话提醒制度，发现问题早提醒、早纠正。')
    add_paragraph_zh(doc, '3. 创新教育形式：利用信息技术优势，探索"微课堂""案例视频"等新形式，增强廉政教育实效。')
    
    # 落款
    doc.add_paragraph()
    doc.add_paragraph()
    
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = closing.add_run('以上报告，请审阅。')
    set_chinese_font(run, '仿宋_GB2312', 12)
    
    # 保存
    if output_path is None:
        output_path = f"/root/.openclaw/workspace/energy_storage/data/述责述廉报告_{datetime.now().strftime('%Y%m%d')}.docx"
    
    doc.save(output_path)
    print(f"报告已生成: {output_path}")
    return output_path

if __name__ == "__main__":
    generate_report()
